import json
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from fastapi import HTTPException
from fastapi.responses import FileResponse
from lxml import etree
from sqlmodel import Session

from app.core.config import EXPORTS_DIR
from app.models import Lecture
from app.services.math_utils import latex_to_text

# ─── OMML namespace ──────────────────────────────────────────────────────────
_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

# ─── Symbol tables (mirrors math_utils.py) ───────────────────────────────────
_GREEK_OMML = {
    'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ',
    'epsilon': 'ε', 'varepsilon': 'ε', 'zeta': 'ζ', 'eta': 'η',
    'theta': 'θ', 'vartheta': 'θ', 'iota': 'ι', 'kappa': 'κ',
    'lambda': 'λ', 'mu': 'μ', 'nu': 'ν', 'xi': 'ξ',
    'pi': 'π', 'varpi': 'π', 'rho': 'ρ', 'varrho': 'ρ',
    'sigma': 'σ', 'varsigma': 'ς', 'tau': 'τ', 'upsilon': 'υ',
    'phi': 'φ', 'varphi': 'φ', 'chi': 'χ', 'psi': 'ψ', 'omega': 'ω',
    'Gamma': 'Γ', 'Delta': 'Δ', 'Theta': 'Θ', 'Lambda': 'Λ',
    'Xi': 'Ξ', 'Pi': 'Π', 'Sigma': 'Σ', 'Upsilon': 'Υ',
    'Phi': 'Φ', 'Psi': 'Ψ', 'Omega': 'Ω',
}
_SYM_OMML = {
    'cdot': '·', 'times': '×', 'div': '÷', 'pm': '±', 'mp': '∓',
    'infty': '∞', 'partial': '∂', 'nabla': '∇', 'prime': '′',
    'sum': '∑', 'prod': '∏', 'int': '∫', 'oint': '∮',
    'leq': '≤', 'le': '≤', 'geq': '≥', 'ge': '≥',
    'neq': '≠', 'ne': '≠', 'approx': '≈', 'equiv': '≡', 'sim': '~',
    'rightarrow': '→', 'to': '→', 'leftarrow': '←', 'mapsto': '↦',
    'Rightarrow': '⇒', 'Leftarrow': '⇐', 'Leftrightarrow': '⟺',
    'in': '∈', 'notin': '∉', 'subset': '⊂', 'supset': '⊃',
    'cup': '∪', 'cap': '∩', 'emptyset': '∅',
    'forall': '∀', 'exists': '∃',
    'circ': '∘', 'propto': '∝',
    'ldots': '…', 'cdots': '⋯', 'vdots': '⋮',
}

# ─── OMML element helpers ────────────────────────────────────────────────────

def _mel(tag: str, *children) -> etree._Element:
    """Create an OMML element with optional children."""
    el = etree.Element(f'{{{_M}}}{tag}')
    for c in children:
        if isinstance(c, list):
            for item in c:
                if item is not None:
                    el.append(item)
        elif c is not None:
            el.append(c)
    return el


def _mr(text: str) -> etree._Element:
    """Create an OMML math-run element."""
    r = _mel('r')
    t = _mel('t')
    t.text = str(text)
    if t.text and (' ' in t.text or (t.text[:1] == ' ') or (t.text[-1:] == ' ')):
        t.set(_XML_SPACE, 'preserve')
    r.append(t)
    return r


def _get_group(s: str, i: int) -> tuple[str, int]:
    """
    Extract a {group} or single token starting at position i.
    Returns (content, position_after).
    """
    while i < len(s) and s[i] == ' ':
        i += 1
    if i >= len(s):
        return '', i
    if s[i] == '{':
        depth, start, i = 0, i + 1, i + 1
        while i < len(s):
            if s[i] == '{':
                depth += 1
            elif s[i] == '}':
                if depth == 0:
                    return s[start:i], i + 1
                depth -= 1
            i += 1
        return s[start:], i
    # Single token: command or character
    if s[i] == '\\':
        m = re.match(r'\\([a-zA-Z]+|[^a-zA-Z\s])', s[i:])
        if m:
            return m.group(0), i + len(m.group(0))
    return s[i], i + 1


def _build_nodes(latex: str) -> list:
    """
    Recursively build a list of OMML elements from a LaTeX math string.
    Handles: fractions, square roots, subscripts, superscripts, Greek
    letters, operators, and plain text runs.
    """
    nodes: list = []
    i = 0
    buf = ''

    def flush():
        nonlocal buf
        if buf:
            nodes.append(_mr(buf))
            buf = ''

    while i < len(latex):
        c = latex[i]

        # ── LaTeX command ────────────────────────────────────────────────────
        if c == '\\':
            m = re.match(r'\\([a-zA-Z]+|[^a-zA-Z\s])', latex[i:])
            if not m:
                i += 1
                continue
            cmd = m.group(1)
            i += len(m.group(0))

            if cmd == 'frac':
                flush()
                num_s, i = _get_group(latex, i)
                den_s, i = _get_group(latex, i)
                f = _mel('f')
                num = _mel('num')
                for n in _build_nodes(num_s):
                    num.append(n)
                den = _mel('den')
                for n in _build_nodes(den_s):
                    den.append(n)
                f.append(num)
                f.append(den)
                nodes.append(f)

            elif cmd == 'sqrt':
                flush()
                inner_s, i = _get_group(latex, i)
                rad = _mel('rad')
                rpr = _mel('radPr')
                dh = _mel('degHide')
                dh.set(f'{{{_M}}}val', '1')
                rpr.append(dh)
                e = _mel('e')
                for n in _build_nodes(inner_s):
                    e.append(n)
                rad.append(rpr)
                rad.append(_mel('deg'))   # empty degree = square root
                rad.append(e)
                nodes.append(rad)

            elif cmd in _GREEK_OMML:
                flush()
                nodes.append(_mr(_GREEK_OMML[cmd]))

            elif cmd in _SYM_OMML:
                flush()
                nodes.append(_mr(' ' + _SYM_OMML[cmd] + ' '))

            elif cmd in ('text', 'mathrm', 'mathbf', 'mathit',
                         'mathbb', 'mathcal', 'operatorname'):
                inner_s, i = _get_group(latex, i)
                buf += inner_s

            elif cmd in ('left', 'right'):
                # Consume optional bracket character
                while i < len(latex) and latex[i] == ' ':
                    i += 1
                if i < len(latex) and latex[i] not in ('\\', '{', ' '):
                    ch = latex[i]
                    i += 1
                    if ch != '.':
                        buf += ch

            elif cmd in ('displaystyle', 'textstyle', 'scriptstyle',
                         'normalsize', 'limits', 'nolimits'):
                pass  # discard

            elif cmd in (',', ';'):
                buf += ' '

            elif cmd in ('quad', 'qquad'):
                buf += '  '

            elif cmd == '\\':
                buf += ' '

            elif cmd == ' ':
                buf += ' '

            else:
                # Unknown command — pass through as text
                buf += cmd

            continue

        # ── Opening brace ────────────────────────────────────────────────────
        if c == '{':
            inner_s, i = _get_group(latex, i)
            flush()
            for n in _build_nodes(inner_s):
                nodes.append(n)
            continue

        # ── Closing brace (stray) ────────────────────────────────────────────
        if c == '}':
            i += 1
            continue

        # ── Subscript / superscript ──────────────────────────────────────────
        if c in ('_', '^'):
            i += 1
            arg_s, i = _get_group(latex, i)
            arg_nodes = _build_nodes(arg_s)
            flush()

            base = [nodes.pop()] if nodes else [_mr('')]

            # Peek for paired operator (a_{x}^{y} or a^{y}_{x})
            j = i
            while j < len(latex) and latex[j] == ' ':
                j += 1
            paired_op = latex[j] if j < len(latex) and latex[j] in ('_', '^') and latex[j] != c else None

            if paired_op:
                i = j + 1
                paired_s, i = _get_group(latex, i)
                paired_nodes = _build_nodes(paired_s)

                sss = _mel('sSubSup')
                e = _mel('e')
                for n in base:
                    e.append(n)
                sub_nodes = arg_nodes if c == '_' else paired_nodes
                sup_nodes = paired_nodes if c == '_' else arg_nodes
                sub_el = _mel('sub')
                for n in sub_nodes:
                    sub_el.append(n)
                sup_el = _mel('sup')
                for n in sup_nodes:
                    sup_el.append(n)
                sss.append(e)
                sss.append(sub_el)
                sss.append(sup_el)
                nodes.append(sss)
            else:
                tag = 'sSub' if c == '_' else 'sSup'
                inner_tag = 'sub' if c == '_' else 'sup'
                ss = _mel(tag)
                e = _mel('e')
                for n in base:
                    e.append(n)
                inner = _mel(inner_tag)
                for n in arg_nodes:
                    inner.append(n)
                ss.append(e)
                ss.append(inner)
                nodes.append(ss)
            continue

        # ── Regular character ────────────────────────────────────────────────
        buf += c
        i += 1

    flush()
    return nodes


def _build_omml(latex: str) -> etree._Element:
    """
    Build an <m:oMath> element from a LaTeX math expression.
    Strips surrounding delimiters ($$, $, \\[, \\() automatically.
    """
    latex = latex.strip()
    for start, end in [('$$', '$$'), (r'\[', r'\]'), ('$', '$'), (r'\(', r'\)')]:
        if latex.startswith(start) and latex.endswith(end) and len(latex) > len(start) + len(end):
            latex = latex[len(start):-len(end)].strip()
            break

    omath = etree.Element(f'{{{_M}}}oMath')
    for node in _build_nodes(latex):
        omath.append(node)
    return omath


# ─── Math-segment splitter ────────────────────────────────────────────────────

_DISPLAY_RE = re.compile(r'(\$\$.+?\$\$|\\\[.+?\\\])', re.DOTALL)
_INLINE_RE  = re.compile(r'(\$.+?\$|\\\(.+?\\\))',      re.DOTALL)


def _split_segments(text: str) -> list[tuple[str, str]]:
    """
    Split text into segments tagged as 'display', 'inline', or 'text'.
    Returns list of (tag, content) tuples.
    """
    segments: list[tuple[str, str]] = []
    for part in _DISPLAY_RE.split(text):
        if _DISPLAY_RE.fullmatch(part):
            # Strip delimiters
            inner = re.sub(r'^\$\$|\$\$$|^\\\[|\\\]$', '', part).strip()
            segments.append(('display', inner))
        else:
            for sub in _INLINE_RE.split(part):
                if _INLINE_RE.fullmatch(sub):
                    inner = re.sub(r'^\$|\$$|^\\\(|\\\)$', '', sub).strip()
                    segments.append(('inline', inner))
                elif sub:
                    segments.append(('text', sub))
    return segments


# ─── python-docx helpers ──────────────────────────────────────────────────────

def _load_processed_content(lecture: Lecture) -> dict:
    try:
        return json.loads(lecture.processed_content_json)
    except Exception:
        return {"topics": [], "summaries": [], "flashcards": []}


def _safe_title(value: str, fallback: str) -> str:
    safe = "".join(c for c in value if c not in r'\/:*?"<>|').strip() or fallback
    return safe.replace('\n', ' ').replace('\r', ' ')[:200]


def _set_paragraph_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1')


def _set_run_rtl(run) -> None:
    rPr = run._element.get_or_add_rPr()
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)
    rtl.set(qn('w:val'), '1')
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ('ascii', 'hAnsi', 'cs'):
        rFonts.set(qn(f'w:{attr}'), 'Arial')


def _add_display_math_paragraph(document: Document, latex: str) -> None:
    """Add a centred display-math paragraph using OMML."""
    try:
        omath = _build_omml(latex)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p._p.append(omath)
    except Exception:
        # Fallback to Unicode text
        _add_rtl_paragraph(document, latex_to_text(f'$${latex}$$'),
                           font_size=12)


def _add_rtl_paragraph(document: Document, text: str, *,
                       font_size: int, bold: bool = False) -> None:
    """
    Add an RTL paragraph that handles mixed Hebrew text and LaTeX math.

    • Display math ($$...$$) becomes a centred OMML paragraph.
    • Inline math ($...$) is embedded as an OMML run inside the paragraph.
    • Plain Hebrew/English text is added as RTL runs.
    """
    segments = _split_segments(text)
    has_display = any(tag == 'display' for tag, _ in segments)

    if has_display:
        # Display math must live in its own paragraph — flush mixed segments
        # into individual paragraphs.
        for tag, content in segments:
            if tag == 'display':
                _add_display_math_paragraph(document, content)
            elif content.strip():
                _add_rtl_paragraph(document, content,
                                   font_size=font_size, bold=bold)
        return

    # No display math — build a single RTL paragraph with optional inline OMML.
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_rtl(paragraph)

    for tag, content in segments:
        if tag == 'inline':
            # Embed OMML directly in the paragraph run
            try:
                omath = _build_omml(content)
                paragraph._p.append(omath)
            except Exception:
                run = paragraph.add_run(latex_to_text(f'${content}$'))
                run.font.size = Pt(font_size)
                run.bold = bold
                run.font.name = 'Arial'
                _set_run_rtl(run)
        else:
            # Plain text run (already decoded; no raw LaTeX expected here)
            if content:
                run = paragraph.add_run(content)
                run.bold = bold
                run.font.size = Pt(font_size)
                run.font.name = 'Arial'
                _set_run_rtl(run)


# ─── Public export function ───────────────────────────────────────────────────

def export_lecture_docx(lecture_id: int, session: Session):
    lecture = session.get(Lecture, lecture_id)
    if not lecture or lecture.status != 'completed':
        raise HTTPException(status_code=404, detail='Lecture not ready')

    data = _load_processed_content(lecture)
    document = Document()
    document.core_properties.title = lecture.title

    # ── Title ────────────────────────────────────────────────────────────────
    _add_rtl_paragraph(document, f'סיכום הרצאה: {lecture.title}',
                       font_size=20, bold=True)

    # ── Key topics ───────────────────────────────────────────────────────────
    if data.get('topics'):
        _add_rtl_paragraph(document, 'נושאים מרכזיים', font_size=16, bold=True)
        for topic in data['topics']:
            _add_rtl_paragraph(document, f'• {topic}', font_size=12)

    # ── Summaries ────────────────────────────────────────────────────────────
    if data.get('summaries'):
        _add_rtl_paragraph(document, 'סיכום מורחב', font_size=16, bold=True)
        for item in data['summaries']:
            _add_rtl_paragraph(document, item.get('topic_name', 'נושא'),
                               font_size=13, bold=True)
            _add_rtl_paragraph(document, item.get('content', ''), font_size=12)

    # ── Flashcards ───────────────────────────────────────────────────────────
    if data.get('flashcards'):
        _add_rtl_paragraph(document, 'כרטיסיות זיכרון', font_size=16, bold=True)
        for idx, card in enumerate(data['flashcards'], 1):
            _add_rtl_paragraph(
                document,
                f'{idx}. שאלה: {card.get("question", "")}',
                font_size=12, bold=True,
            )
            _add_rtl_paragraph(
                document,
                f'תשובה: {card.get("answer", "")}',
                font_size=12,
            )

    safe_title = _safe_title(lecture.title, f'lecture-{lecture_id}')
    export_path = EXPORTS_DIR / f'export_{lecture_id}.docx'
    document.save(str(export_path))

    return FileResponse(
        export_path,
        filename=f'summary-{safe_title}.docx',
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
